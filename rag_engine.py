"""
RAG Engine — Core pipeline for document processing and question answering.
Handles: PDF/TXT/DOCX/CSV loading, text chunking, embedding, vector storage, and retrieval-augmented generation.
"""

import os
import csv
import io
from typing import List, Tuple

from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import FAISS
from langchain.schema import Document, HumanMessage, SystemMessage

from config import (
    CHUNK_SIZE,
    CHUNK_OVERLAP,
    TOP_K_RESULTS,
    EMBEDDING_MODEL,
    CHAT_MODEL,
    TEMPERATURE,
    SYSTEM_PROMPT,
)


# ──────────────────────────────────────────────
# 1. Document Loading (Multi-Format)
# ──────────────────────────────────────────────

def load_pdf(uploaded_file) -> str:
    """Extract text from an uploaded PDF file."""
    reader = PdfReader(uploaded_file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text


def load_txt(uploaded_file) -> str:
    """Extract text from an uploaded TXT file."""
    content = uploaded_file.read()
    if isinstance(content, bytes):
        content = content.decode("utf-8", errors="ignore")
    return content


def load_docx(uploaded_file) -> str:
    """Extract text from an uploaded DOCX (Word) file."""
    doc = DocxDocument(uploaded_file)
    text = ""
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            text += row_text + "\n"
    return text


def load_csv(uploaded_file) -> str:
    """Extract text from an uploaded CSV file."""
    content = uploaded_file.read()
    if isinstance(content, bytes):
        content = content.decode("utf-8", errors="ignore")
    reader = csv.reader(io.StringIO(content))
    text = ""
    for row in reader:
        text += " | ".join(row) + "\n"
    return text


def load_document(uploaded_file) -> str:
    """Load a document based on its file extension.
    
    Supports: PDF, TXT, DOCX, CSV
    """
    filename = uploaded_file.name.lower()
    
    if filename.endswith(".pdf"):
        return load_pdf(uploaded_file)
    elif filename.endswith(".txt"):
        return load_txt(uploaded_file)
    elif filename.endswith(".docx"):
        return load_docx(uploaded_file)
    elif filename.endswith(".csv"):
        return load_csv(uploaded_file)
    else:
        raise ValueError(f"Unsupported file format: {filename}")


def load_multiple_documents(uploaded_files) -> List[dict]:
    """Extract text from multiple uploaded files (any supported format).
    
    Returns:
        List of dicts with 'filename' and 'text' keys.
    """
    documents = []
    for file in uploaded_files:
        try:
            text = load_document(file)
            if text.strip():
                documents.append({
                    "filename": file.name,
                    "text": text
                })
        except Exception as e:
            # Skip files that fail to load, errors shown in UI
            print(f"Error loading {file.name}: {e}")
    return documents


# ──────────────────────────────────────────────
# 2. Text Chunking
# ──────────────────────────────────────────────

def split_into_chunks(text: str, filename: str = "document") -> List[Document]:
    """Split text into overlapping chunks for embedding.
    
    Args:
        text: The full document text.
        filename: Source filename for metadata.
    
    Returns:
        List of LangChain Document objects with metadata.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    
    chunks = splitter.split_text(text)
    
    documents = [
        Document(
            page_content=chunk,
            metadata={"source": filename, "chunk_index": i}
        )
        for i, chunk in enumerate(chunks)
    ]
    
    return documents


# ──────────────────────────────────────────────
# 3. Vector Store (Embedding + Storage)
# ──────────────────────────────────────────────

def create_vector_store(documents: List[Document], api_key: str) -> FAISS:
    """Create a FAISS vector store from document chunks.
    
    Args:
        documents: List of LangChain Document objects.
        api_key: OpenAI API key for generating embeddings.
    
    Returns:
        FAISS vector store instance.
    """
    embeddings = OpenAIEmbeddings(
        model=EMBEDDING_MODEL,
        openai_api_key=api_key
    )
    
    vector_store = FAISS.from_documents(documents, embeddings)
    return vector_store


def add_to_vector_store(vector_store: FAISS, documents: List[Document], api_key: str) -> FAISS:
    """Add new documents to an existing vector store.
    
    Args:
        vector_store: Existing FAISS vector store.
        documents: New documents to add.
        api_key: OpenAI API key.
    
    Returns:
        Updated FAISS vector store.
    """
    embeddings = OpenAIEmbeddings(
        model=EMBEDDING_MODEL,
        openai_api_key=api_key
    )
    
    vector_store.add_documents(documents)
    return vector_store


# ──────────────────────────────────────────────
# 4. Retrieval
# ──────────────────────────────────────────────

def get_relevant_context(query: str, vector_store: FAISS, k: int = TOP_K_RESULTS) -> List[Tuple[Document, float]]:
    """Retrieve the most relevant document chunks for a query.
    
    Args:
        query: User's question.
        vector_store: FAISS vector store to search.
        k: Number of results to return.
    
    Returns:
        List of (Document, score) tuples, sorted by relevance.
    """
    results = vector_store.similarity_search_with_score(query, k=k)
    return results


# ──────────────────────────────────────────────
# 5. Question Answering (RAG)
# ──────────────────────────────────────────────

def format_context(results: List[Tuple[Document, float]]) -> str:
    """Format retrieved documents into a context string.
    
    Args:
        results: List of (Document, score) tuples.
    
    Returns:
        Formatted context string with source references.
    """
    context_parts = []
    for i, (doc, score) in enumerate(results, 1):
        source = doc.metadata.get("source", "Unknown")
        context_parts.append(
            f"[Source: {source} | Chunk {doc.metadata.get('chunk_index', '?')}]\n{doc.page_content}"
        )
    
    return "\n\n---\n\n".join(context_parts)


def ask_question(
    query: str,
    vector_store: FAISS,
    chat_history: List[dict],
    api_key: str
) -> Tuple[str, List[Tuple[Document, float]]]:
    """Answer a question using RAG (Retrieval-Augmented Generation).
    
    Args:
        query: User's question.
        vector_store: FAISS vector store with embedded documents.
        chat_history: Previous conversation messages.
        api_key: OpenAI API key.
    
    Returns:
        Tuple of (answer_text, source_documents).
    """
    # Step 1: Retrieve relevant context
    relevant_docs = get_relevant_context(query, vector_store)
    context = format_context(relevant_docs)
    
    # Step 2: Build the prompt with context
    augmented_prompt = f"""Based on the following context from the uploaded document(s), answer the user's question.

--- DOCUMENT CONTEXT ---
{context}
--- END CONTEXT ---

User's Question: {query}

Instructions: Answer the question using ONLY the information from the context above. If the answer is not in the context, say so clearly."""

    # Step 3: Build messages list with conversation history
    messages = [SystemMessage(content=SYSTEM_PROMPT)]
    
    # Add recent chat history (last 6 exchanges to stay within token limits)
    recent_history = chat_history[-12:]  # 12 messages = 6 Q&A pairs
    for msg in recent_history:
        if msg["role"] == "user":
            messages.append(HumanMessage(content=msg["content"]))
        else:
            from langchain.schema import AIMessage
            messages.append(AIMessage(content=msg["content"]))
    
    # Add current augmented question
    messages.append(HumanMessage(content=augmented_prompt))
    
    # Step 4: Call OpenAI
    llm = ChatOpenAI(
        model=CHAT_MODEL,
        temperature=TEMPERATURE,
        openai_api_key=api_key,
        streaming=True
    )
    
    response = llm.invoke(messages)
    
    return response.content, relevant_docs


def get_streaming_response(
    query: str,
    vector_store: FAISS,
    chat_history: List[dict],
    api_key: str
):
    """Stream a response token-by-token using RAG.
    
    Args:
        query: User's question.
        vector_store: FAISS vector store.
        chat_history: Previous conversation messages.
        api_key: OpenAI API key.
    
    Yields:
        Tuples of (token, source_docs). source_docs is None until the last yield.
    """
    # Retrieve relevant context
    relevant_docs = get_relevant_context(query, vector_store)
    context = format_context(relevant_docs)
    
    augmented_prompt = f"""Based on the following context from the uploaded document(s), answer the user's question.

--- DOCUMENT CONTEXT ---
{context}
--- END CONTEXT ---

User's Question: {query}

Instructions: Answer the question using ONLY the information from the context above. If the answer is not in the context, say so clearly."""

    # Build messages
    messages = [SystemMessage(content=SYSTEM_PROMPT)]
    
    recent_history = chat_history[-12:]
    for msg in recent_history:
        if msg["role"] == "user":
            messages.append(HumanMessage(content=msg["content"]))
        else:
            from langchain.schema import AIMessage
            messages.append(AIMessage(content=msg["content"]))
    
    messages.append(HumanMessage(content=augmented_prompt))
    
    # Stream from OpenAI
    llm = ChatOpenAI(
        model=CHAT_MODEL,
        temperature=TEMPERATURE,
        openai_api_key=api_key,
        streaming=True
    )
    
    for chunk in llm.stream(messages):
        if chunk.content:
            yield chunk.content, relevant_docs
